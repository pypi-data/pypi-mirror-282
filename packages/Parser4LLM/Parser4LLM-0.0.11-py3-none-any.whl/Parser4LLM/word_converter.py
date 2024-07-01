import os
from docx.document import Document as DocumentObject
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.run import Run
from typing import List
from collections import Counter
from Parser4LLM.notification import Notification, DefaultNotification
    
class DOCXConverter():
    def __init__(self, notification: Notification = DefaultNotification()) -> None:
        self.header_handled = False
        self.notification = notification
        
    def is_ocr_docx(self, file_path):
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                return True  # OCR if any paragraph has text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        return True  # OCR if any table cell has text
        # If no text could be extracted from paragraphs or tables, it's likely non-OCR
        return False

    def convert(self, file_path: str) -> str:
        is_ocr = self.is_ocr_docx(file_path)
        if is_ocr:
            doc = Document(file_path)
            md_content = []
            # Handle header
            if doc.sections and doc.sections[0].header:
                header_content = self._handle_header(doc.sections[0].header)
                if header_content:
                    md_content.append(header_content)

            for element in doc.element.body:
                if isinstance(element, CT_P):
                    md_content.append(self._handle_paragraph(Paragraph(element, doc)))
                elif isinstance(element, CT_Tbl):
                    md_content += self._handle_table(Table(element, doc))
                # Add more handlers here (image, header, footer, etc)

            return "\n".join(md_content)
        else:
            # call marker-api
            return "No Extracted text"

    def _handle_header(self, header) -> str:
        if not self.header_handled:
            parts = []
            for paragraph in header.paragraphs:
                parts.append(f"# {paragraph.text}")
            for table in header.tables:
                parts += self._handle_header_table(table)
            self.header_handled = True
            return "\n".join(parts)
        return ""
    
    def _handle_header_table(self, table: Table) -> List[str]:
        cell_texts = [cell.text for row in table.rows for cell in row.cells]
        if "" in cell_texts:
            cell_texts.remove("")
        # Find the most repeated cell text
        text_counts = Counter(cell_texts)
        title = text_counts.most_common(1)[0][0] if cell_texts else ""
        other_texts = [text for text in cell_texts if text != title and text != ""]

        md_table_content = []
        if title:
            md_table_content.append(f"# {title}")
        for text in other_texts:
            md_table_content.append(f"*{text}*;")
        return md_table_content

    def _handle_paragraph(self, paragraph) -> str:
        if paragraph.style.name.startswith("Heading"):  # type: ignore
            level = int(paragraph.style.name.split()[-1])  # type: ignore
            return f"{'#' * level} {paragraph.text}"
        else:
            parts = []
            for run in paragraph.runs:
                if run.text != "":
                    parts.append(self._handle_run(run))
            return "".join(parts)

    def _handle_run(self, run) -> str:
        text: str = run.text
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        return text

    def _handle_table(self, table):
        row_content = []
        for i, row in enumerate(table.rows):
            row_content.append(
                "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |"
            )
            if i == 0:
                row_content.append("|" + "---|" * len(row.cells))

        return row_content

    def save_md(self, md_content: str, file_path: str, encoding: str = "utf-8") -> None:
        with open(file_path, "w", encoding=encoding) as f:
            f.write(md_content)