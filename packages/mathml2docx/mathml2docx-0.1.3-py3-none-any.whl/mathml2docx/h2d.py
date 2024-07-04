"""
Make 'span' in tags dict a stack
maybe do the same for all tags in case of unclosed tags?
optionally use bs4 to clean up invalid html?

the idea is that there is a method that converts html files into docx
but also have api methods that let user have more control e.g. so they
can nest calls to something like 'convert_chunk' in loops

user can pass existing document object as arg 
(if they want to manage rest of document themselves)

How to deal with block level style applied over table elements? e.g. text align
"""
import re, argparse
import io, os, base64
import urllib.request
from html.parser import HTMLParser

import docx, docx.table
from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_COLOR, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from bs4 import BeautifulSoup
from lxml import etree

from .config import *
from .utils import *

class HtmlToDocx(HTMLParser):
    def __init__(self):
        super().__init__()
        self.options = {
            'fix-html': True,
            'images': True,
            'tables': True,
            'mathml': True,
            'styles': True,
            'img_src_base_path':None,
        }
        self.table_row_selectors = [
            'table > tr',
            'table > thead > tr',
            'table > tbody > tr',
            'table > tfoot > tr'
        ]
        self.table_style = DEFAULT_TABLE_STYLE
        self.paragraph_style = DEFAULT_PARAGRAPH_STYLE

    def set_initial_attrs(self, document=None):
        self.tags = {
            'span': [],
            'list': [],
        }
        if document:
            self.doc = document
        else:
            self.doc = Document()
        self.bs = self.options['fix-html'] # whether or not to clean with BeautifulSoup
        self.document = self.doc
        self.include_tables = self.options['tables']
        self.include_images = self.options['images']
        self.include_mathml = self.options['mathml']
        self.include_styles = self.options['styles']
        self.img_src_base_path = self.options['img_src_base_path']
        self.paragraph = None
        self.run = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0

    def copy_settings_from(self, other):
        """Copy settings from another instance of HtmlToDocx"""
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style

    def get_cell_html(self, soup):
        # Returns string of td element with opening and closing <td> tags removed
        # Cannot use find_all as it only finds element tags and does not find text which
        # is not inside an element
        return ' '.join([str(i) for i in soup.contents])

    def add_styles_to_paragraph(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style and 'margin-right' in style:
            margin_left = style['margin-left']
            margin_right = style['margin-right']
            if "auto" in margin_left and "auto" in margin_right:
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin_suffix = re.sub(r'[a-z]+', '', margin)
            if len(margin_suffix) > 0:
                margin = int(float(margin_suffix))
                if units == 'px':
                    self.paragraph.paragraph_format.left_indent = Inches(min(margin // 10 * INDENT, MAX_INDENT))
                # TODO handle non px units

    def add_styles_to_table(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.table.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style and 'margin-right' in style:
            margin_left = style['margin-left']
            margin_right = style['margin-right']
            if "auto" in margin_left and "auto" in margin_right:
                self.table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin_suffix = re.sub(r'[a-z]+', '', margin)
            if len(margin_suffix) > 0:
                margin = int(float(margin_suffix))
                if units == 'px':
                    self.table.left_indent = Inches(min(margin // 10 * INDENT, MAX_INDENT))
                # TODO handle non px units

    def add_styles_to_run(self, style):
        if 'color' in style:
            if 'rgb' in style['color']:
                color = re.sub(r'[a-z()]+', '', style['color'])
                colors = [int(x) for x in color.split(',')][:3]
            elif '#' in style['color']:
                color = style['color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
                # TODO map colors to named colors (and extended colors...)
                # For now set color to black to prevent crashing

            self.run.font.color.rgb = RGBColor(*colors)
            
        if 'background-color' in style:
            if 'rgb' in style['background-color']:
                color = color = re.sub(r'[a-z()]+', '', style['background-color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['background-color']:
                color = style['background-color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
                # TODO map colors to named colors (and extended colors...)
                # For now set color to black to prevent crashing
            self.run.font.highlight_color = WD_COLOR.GRAY_25 #TODO: map colors

    def apply_paragraph_style(self, style=None):
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as e:
            raise ValueError(f"Unable to apply style {self.paragraph_style}.") from e

    def parse_dict_string(self, string, separator=';'):
        new_string = string.replace(" ", '').split(separator)
        string_dict = dict([x.split(':') for x in new_string if ':' in x])
        return string_dict

    def handle_li(self):
        # check list stack to determine style and depth
        list_depth = len(self.tags['list'])
        if list_depth:
            list_type = self.tags['list'][-1]
        else:
            list_type = 'ul' # assign unordered if no tag

        if list_type == 'ol':
            list_style = styles['LIST_NUMBER']
        else:
            list_style = styles['LIST_BULLET']

        self.paragraph = self.doc.add_paragraph(style=list_style)            
        self.paragraph.paragraph_format.left_indent = Inches(min(list_depth * LIST_INDENT, MAX_INDENT))
        self.paragraph.paragraph_format.line_spacing = 1

    def add_image_to_cell(self, cell, image):
        # python-docx doesn't have method yet for adding images to table cells. For now we use this
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image)

    def handle_img(self, attrs):
        if not self.include_images:
            self.skip = True
            self.skip_tag = 'img'
            return
        
        src_attr = attrs.get('src')

        if not src_attr:
            if self.paragraph:
                self.paragraph.add_run("<image: no_src>")
            else:
                self.doc.add_paragraph("<image: no_src>")
            return
        
        # fetch image
        src_is_url = is_url(src_attr)
        if src_is_url:
            try:
                image = fetch_image(src_attr)
            except urllib.error.URLError:
                image = None
        else:
            
            image = src_attr
            print(self.img_src_base_path)
            if image and image.startswith('data:image/'):
                #-- convert to bytes ready to insert to docx
                image = image.split(',')[1]
                image = base64.b64decode(image)
                image = io.BytesIO(image)
            elif image and self.img_src_base_path:
                image = os.path.abspath(os.path.join(self.img_src_base_path, image))
                print(image)
                if os.path.isfile(image):
                    with open(image, "rb") as fp:
                        image = io.BytesIO(fp.read())
                else:
                    image = None
            else:
                image = None
                
        # add image to doc
        if image:
            try:
                if isinstance(self.doc, docx.document.Document):
                    if self.paragraph:
                        self.run = self.paragraph.add_run()
                        self.run.add_picture(image)
                    else:
                        self.doc.add_picture(image)
                else:
                    self.add_image_to_cell(self.doc, image)
            except FileNotFoundError:
                image = None
        else:
            # avoid exposing filepaths in document
            msg = "<image: %s>" % src_attr if src_is_url else "<image: %s>" % get_filename_from_url(src_attr)

            if self.paragraph:
                self.paragraph.add_run(msg)
            else:
                self.doc.add_paragraph(msg)
                
        # add styles?

    def handle_table(self, current_attrs):
        """
        To handle nested tables, we will parse tables manually as follows:
        Get table soup
        Create docx table
        Iterate over soup and fill docx table with new instances of this parser
        Tell HTMLParser to ignore any tags until the corresponding closing table tag
        """
        if not self.include_tables:
            self.skip = True
            self.skip_tag = 'table'
            return
        
        table_soup = self.tables[self.table_no]
        rows, cols = self.get_table_dimensions(table_soup)
        self.table = self.doc.add_table(rows, cols)

        if self.table_style:
            try:
                self.table.style = self.table_style
            except KeyError as e:
                raise ValueError(f"Unable to apply style {self.table_style}.") from e

        rows = self.get_table_rows(table_soup)
        cell_row = 0
        for row in rows:
            cols = self.get_table_columns(row)
            cell_col = 0
            for col in cols:
                colspan = int(col.attrs.get('colspan', 1))
                rowspan = int(col.attrs.get('rowspan', 1))

                cell_html = self.get_cell_html(col)
                if col.name == 'th':
                    cell_html = "<b>%s</b>" % cell_html

                docx_cell = self.table.cell(cell_row, cell_col)
                while docx_cell.text != '':  # Skip the merged cell
                    cell_col += 1
                    docx_cell = self.table.cell(cell_row, cell_col)

                cell_to_merge = self.table.cell(cell_row + rowspan - 1, cell_col + colspan - 1)
                if docx_cell != cell_to_merge:
                    docx_cell.merge(cell_to_merge)

                child_parser = HtmlToDocx()
                child_parser.copy_settings_from(self)
                child_parser.add_html_to_cell(cell_html or ' ', docx_cell)  # occupy the position

                cell_col += colspan
            cell_row += 1
        
        if 'style' in current_attrs and self.table:
            style = self.parse_dict_string(current_attrs['style'])
            self.add_styles_to_table(style)
        
        # skip all tags until corresponding closing tag
        self.instances_to_skip = len(table_soup.find_all('table'))
        self.skip_tag = 'table'
        self.skip = True
        self.table = None
    
    def handle_mathml(self, attr):
        if not self.include_mathml:
            self.skip = True
            self.skip_tag = 'math'
            return
        
        math_soup = self.mathml_list[self.math_no]
        mathml_string = str(math_soup)
        
        tree = etree.fromstring(mathml_string)
        xslt = etree.parse(MML2OMML_PATH)

        transform = etree.XSLT(xslt)
        new_dom = transform(tree)

        display_attr = attr.get("display", "block")
        
        is_require_new_para = False
        if self.paragraph:
            if display_attr != "inline":
                is_require_new_para = True
        else:
            is_require_new_para = True
            
        if is_require_new_para:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()

        self.paragraph._element.append(new_dom.getroot())

        self.instances_to_skip = len(math_soup.find_all('math'))
        self.skip_tag = 'math'
        self.skip = True

    def handle_div(self, current_attrs):
        # handle page break
        if 'style' in current_attrs and "page-break-after: always" in current_attrs['style']:
            self.doc.add_page_break()

    def handle_link(self, href, text):
        # Link requires a relationship
        is_external = href.startswith('http')
        rel_id = self.paragraph.part.relate_to(
            href,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True  # don't support anchor links for this library yet
        )

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), rel_id)


        # Create sub-run
        subrun = self.paragraph.add_run()
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # add default color
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), "0000EE")
        rPr.append(c)

        # add underline
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

        subrun._r.append(rPr)
        subrun._r.text = text

        # Add subrun to hyperlink
        hyperlink.append(subrun._r)

        # Add hyperlink to run
        self.paragraph._p.append(hyperlink)

    def handle_starttag(self, tag, attrs):
        if self.skip:
            return
        if tag == 'head':
            self.skip = True
            self.skip_tag = tag
            self.instances_to_skip = 0
            return
        elif tag == 'body':
            return
        current_attrs = dict(attrs)

        if tag == 'span':
            self.tags['span'].append(current_attrs)
            return
        elif tag == 'ol' or tag == 'ul':
            self.tags['list'].append(tag)
            return # don't apply styles for now
        elif tag == 'br':
            if self.paragraph:
                self.run = self.paragraph.add_run()
            else:
                self.paragraph = self.doc.add_paragraph()
                self.apply_paragraph_style()

            if not self.run:
                self.run = self.paragraph.add_run()

            self.run.add_break()
            return

        self.tags[tag] = current_attrs
        if tag in ['p', 'pre']:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()

        elif tag == 'li':
            self.handle_li()

        elif tag == "hr":

            # This implementation was taken from:
            # https://github.com/python-openxml/python-docx/issues/105#issuecomment-62806373

            self.paragraph = self.doc.add_paragraph()
            pPr = self.paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            pPr.insert_element_before(pBdr,
                'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                'w:pPrChange'
            )
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)

        elif re.match('h[1-9]', tag):
            if isinstance(self.doc, docx.document.Document):
                h_size = int(tag[1])
                self.paragraph = self.doc.add_heading(level=min(h_size, 9))
            else:
                self.paragraph = self.doc.add_paragraph()

        elif tag == 'img':
            self.handle_img(current_attrs)
            return

        elif tag == 'table':
            self.handle_table(current_attrs)
            return
        elif tag == 'math':
            self.handle_mathml(current_attrs)
            return
        elif tag == "div":
            self.handle_div(current_attrs)

        # set new run reference point in case of leading line breaks
        if tag in ['p', 'li', 'pre']:
            self.run = self.paragraph.add_run()

        # add style
        if not self.include_styles:
            return
        if 'style' in current_attrs and self.paragraph:
            style = self.parse_dict_string(current_attrs['style'])
            self.add_styles_to_paragraph(style)
            
    def handle_endtag(self, tag):
        if self.skip:
            if not tag == self.skip_tag:
                return

            if self.instances_to_skip > 0:
                self.instances_to_skip -= 1
                return

            self.skip = False
            self.skip_tag = None

        if tag == 'span':
            if self.tags['span']:
                self.tags['span'].pop()
                return
        elif tag == 'ol' or tag == 'ul':
            remove_last_occurence(self.tags['list'], tag)
            return
        elif tag in ['p', 'pre']:
            self.paragraph = None

        elif tag == 'li':
            self.paragraph = None

        elif tag == "hr":
            self.paragraph = None

        elif re.match('h[1-9]', tag):
            self.paragraph = None

        elif tag == 'table':
            if not self.include_tables:
                return
            self.table_no += 1
            self.table = None
            self.doc = self.document
            self.paragraph = None
        elif tag == 'math':
            if not self.include_mathml:
                return
            self.math_no += 1
            self.doc = self.document

        if tag in self.tags:
            self.tags.pop(tag)
        # maybe set relevant reference to None?

    def handle_data(self, data):
        if self.skip:
            return

        # Only remove white space if we're not in a pre block.
        if 'pre' not in self.tags:
            # remove leading and trailing whitespace in all instances
            data = remove_whitespace(data, True, True)

        # There can only be one nested link in a valid html document
        # You cannot have interactive content in an A tag, this includes links
        # https://html.spec.whatwg.org/#interactive-content
        link = self.tags.get('a')
        if link:
            self.handle_link(link['href'], data)
        else:
            if not data:
                return
            
            if not self.paragraph:
                self.paragraph = self.doc.add_paragraph()
                self.apply_paragraph_style()

            # If there's a link, dont put the data directly in the run
            self.run = self.paragraph.add_run(data)
            spans = self.tags['span']
            for span in spans:
                if 'style' in span:
                    style = self.parse_dict_string(span['style'])
                    self.add_styles_to_run(style)

            # add font style and name
            for tag in self.tags:
                if tag in font_styles:
                    font_style = font_styles[tag]
                    setattr(self.run.font, font_style, True)

                if tag in font_names:
                    font_name = font_names[tag]
                    self.run.font.name = font_name

    def ignore_nested_tables(self, tables_soup):
        """
        Returns array containing only the highest level tables
        Operates on the assumption that bs4 returns child elements immediately after
        the parent element in `find_all`. If this changes in the future, this method will need to be updated

        :return:
        """
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all('table'))
        return new_tables

    def get_table_rows(self, table_soup):
        # If there's a header, body, footer or direct child tr tags, add row dimensions from there
        return table_soup.select(', '.join(self.table_row_selectors), recursive=False)

    def get_table_columns(self, row):
        # Get all columns for the specified row tag.
        return row.find_all(['th', 'td'], recursive=False) if row else []

    def get_table_dimensions(self, table_soup):
        # Get rows for the table
        rows = self.get_table_rows(table_soup)
        # Table is either empty or has non-direct children between table and tr tags
        # Thus the row dimensions and column dimensions are assumed to be 0
        
        cols = max([len(self.get_table_columns(row)) for row in rows]) if rows else 0
        return len(rows), cols

    def get_tables(self):
        if not hasattr(self, 'soup'):
            self.include_tables = False
            return
            # find other way to do it, or require this dependency?
        self.tables = self.ignore_nested_tables(self.soup.find_all('table'))  
        self.table_no = 0

    def get_mathml(self):
        if not hasattr(self, 'soup'):
            self.include_mathml = False
            return
        self.mathml_list = self.soup.find_all('math')
        self.math_no = 0

    def run_process(self, html):
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')
            html = str(self.soup)
        if self.include_tables:
            self.get_tables()
        if self.include_mathml:
            self.get_mathml()
        self.feed(html)

    def add_html_to_document(self, html, document):
        if not isinstance(html, str):
            raise ValueError('First argument needs to be a %s' % str)
        elif not isinstance(document, docx.document.Document) and not isinstance(document, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.document.Document)
        self.set_initial_attrs(document)
        self.run_process(html)

    def add_html_to_cell(self, html, cell):
        if not isinstance(cell, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.table._Cell)
        unwanted_paragraph = cell.paragraphs[0]
        if unwanted_paragraph.text == "":
            delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        # cells must end with a paragraph or will get message about corrupt file
        # https://stackoverflow.com/a/29287121
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')  

    def parse_html_file(self, filename_html, filename_docx=None, encoding='utf-8'):
        with open(filename_html, 'r', encoding=encoding) as infile:
            html = infile.read()
        self.set_initial_attrs()
        self.run_process(html)
        if not filename_docx:
            path, filename = os.path.split(filename_html)
            filename_docx = '%s/new_docx_file_%s' % (path, filename)
        self.doc.save('%s.docx' % filename_docx)
    
    def parse_html_string(self, html):
        self.set_initial_attrs()
        self.run_process(html)
        return self.doc

if __name__=='__main__':
    arg_parser = argparse.ArgumentParser(description='Convert .html file into .docx file with formatting')
    arg_parser.add_argument('filename_html', help='The .html file to be parsed')
    arg_parser.add_argument(
        'filename_docx', 
        nargs='?', 
        help='The name of the .docx file to be saved. Default new_docx_file_[filename_html]', 
        default=None
    )
    arg_parser.add_argument('--bs', action='store_true', 
        help='Attempt to fix html before parsing. Requires bs4. Default True')

    args = vars(arg_parser.parse_args())
    file_html = args.pop('filename_html')
    html_parser = HtmlToDocx()
    html_parser.parse_html_file(file_html, **args)